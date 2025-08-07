import streamlit as st
import torch
import pytesseract
from PIL import Image
from transformers import LayoutLMv3Processor, LayoutLMv3ForTokenClassification
import numpy as np
from docx import Document
import tempfile
import os
import fitz  # PyMuPDF for PDF reading
import google.generativeai as genai
import psycopg2
import io
from urllib.parse import urlparse
from datetime import datetime

# === CONFIG ===
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
DATABASE_URL = os.environ.get("DATABASE_URL")
model = genai.GenerativeModel("models/gemma-3n-e4b-it")

def get_db_connection():
    return psycopg2.connect(DATABASE_URL, sslmode='require')

def create_table_if_not_exists():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS prescriptions (
            id SERIAL PRIMARY KEY,
            uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            filename TEXT,
            raw_text TEXT,
            structured_text TEXT
        )
    """)
    conn.commit()
    cursor.close()
    conn.close()

def save_to_neon_db(filename, raw_text, structured_text):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO prescriptions (filename, raw_text, structured_text)
            VALUES (%s, %s, %s)
        """, (filename, raw_text, structured_text))
        conn.commit()
        cursor.close()
        conn.close()
        st.success("✅ Saved to Neon DB.")
    except Exception as e:
        st.error(f"❌ Failed to save to DB: {e}")

# === OCR Functions ===
def get_ocr_data(image):
    data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
    word_data = []
    for i in range(len(data['text'])):
        if int(data['conf'][i]) > 60 and data['text'][i].strip():
            word = data['text'][i]
            x, y, w, h = data['left'][i], data['top'][i], data['width'][i], data['height'][i]
            word_data.append({'text': word, 'box': [x, y, x + w, y + h], 'top': y, 'left': x})
    return word_data

def group_words_into_lines(word_data, line_threshold=15):
    word_data = sorted(word_data, key=lambda x: x['top'])
    lines, current_line, last_top = [], [], None
    for word in word_data:
        if last_top is None or abs(word['top'] - last_top) <= line_threshold:
            current_line.append(word)
        else:
            lines.append(sorted(current_line, key=lambda x: x['left']))
            current_line = [word]
        last_top = word['top']
    if current_line:
        lines.append(sorted(current_line, key=lambda x: x['left']))
    return lines

# === Streamlit UI ===
st.title("🩺 Handwritten Prescription Structuring")
uploaded_file = st.file_uploader("Upload a Prescription Image or PDF", type=["jpg", "jpeg", "png", "pdf"])

# Create DB table if needed
create_table_if_not_exists()

if uploaded_file is not None:
    file_ext = uploaded_file.name.split('.')[-1].lower()

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as tmp_file:
        tmp_file.write(uploaded_file.read())
        temp_path = tmp_file.name

    if uploaded_file.type == "application/pdf":
        pdf = fitz.open(temp_path)
        pix = pdf[0].get_pixmap(dpi=300)
        image_bytes = pix.tobytes("png")
        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    else:
        image = Image.open(temp_path).convert("RGB")

    # OCR and structuring pipeline
    word_data = get_ocr_data(image)
    lines = group_words_into_lines(word_data)

    # Step 2: Save Raw Lines
    doc = Document()
    doc.add_heading('Raw Prescription Text', level=1)
    for line in lines:
        line_text = " ".join([word['text'] for word in line])
        doc.add_paragraph(line_text)
    raw_docx_path = "raw_prescription.docx"
    doc.save(raw_docx_path)
    st.success("✅ Raw text extracted and saved.")

    # Step 3: Format via Gemini
    raw_text = "\n".join([" ".join([word['text'] for word in line]) for line in lines])
    prompt = f"""
Here is a raw OCR extracted text from a handwritten medical prescription.

Please convert it into a **well-structured document** that includes:
- Hospital Name
- Patient Name
- Age / Gender
- Address
- Date of Visit
- Doctor Name & Degree
- Diagnosis (if found)
- Medications with dosage & timing
- Instructions
- Follow-up Advice (if any)

Here is the raw text:
\"\"\" 
{raw_text} 
\"\"\" 

Return only the well-structured formatted result.
"""
    response = model.generate_content(prompt)
    structured_text = response.text

    # Step 4: Save Structured Output
    structured_doc = Document()
    structured_doc.add_heading('Structured Prescription', level=1)
    for paragraph in structured_text.strip().split("\n\n"):
        structured_doc.add_paragraph(paragraph.strip())
    structured_docx_path = "structured_prescription.docx"
    structured_doc.save(structured_docx_path)
    st.success("✅ Structured text generated and saved.")

    # Step 5: Save to Neon
    save_to_neon_db(uploaded_file.name, raw_text, structured_text)

    # Step 6: Downloads
    st.download_button("📄 Download Raw Prescription", data=open(raw_docx_path, "rb").read(), file_name="raw_prescription.docx")
    st.download_button("📑 Download Structured Prescription", data=open(structured_docx_path, "rb").read(), file_name="structured_prescription.docx")

